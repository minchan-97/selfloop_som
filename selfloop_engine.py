"""
selfloop_engine.py
==================
Streamlit 앱(app_selfloop.py)이 import 하는 통합 엔진.

구성:
  - EmbeddingProvider : tok_emb pkl 있으면 사용, 없으면 결정론적 해시 폴백
  - GrowingSOM        : 자율 성장 SOM (+ INT8 양자화 저장)
  - Crawler           : "이 분야 검색해서 학습" 명령 -> 웹 수집 (네트워크 필요)
  - LLMBridge         : 답변 루프용 LLM 호출 (API 키 필요)
  - SelfLoopState     : 전체 상태 + pkl 저장/불러오기(이전 학습 누적)
  - introspect 지표   : 점유율/QE/어휘다양성/추론경로

네트워크/키가 없는 환경에서도 코퍼스 학습·계측·pkl 은 완전 작동.
크롤링/LLM 은 로컬에서 키 넣으면 작동.
"""

from __future__ import annotations
import os
import re
import time
import pickle
import hashlib
import numpy as np


# ======================================================================
# 임베딩
# ======================================================================
class EmbeddingProvider:
    """
    실제 TinyTransformer tok_emb를 받아 사용. 없거나 형식 불명이면 해시 폴백.

    인식하는 pkl 형식(자동 감지):
      1) {"word2idx": {...}, "tok_emb": ndarray(V, dim)}            # 기본
      2) {"word2idx": {...}, "tok_emb": ndarray, "dim": int}        # to_dict 직렬화형
      3) {"word2idx": {...}, "embeddings"/"emb"/"weight": ndarray}  # 키 이름 변형
      4) {"<단어>": ndarray, ...}                                    # 단어->벡터 딕셔너리
      5) gascore_engine 통째 dict 안에 위 키들이 중첩된 경우도 탐색
    self.load_error 에 실패 사유를 남겨 UI에서 표시할 수 있다.
    """
    EMB_KEYS = ("tok_emb", "embeddings", "embedding", "emb", "weight", "weights", "vectors")

    def __init__(self, dim=64, tok_emb_path: str | None = None):
        self.dim = dim
        self.mode = "hash"
        self.word2idx = None
        self.tok_emb = None
        self.load_error: str | None = None
        self.vocab_size = 0
        self._cache: dict[str, np.ndarray] = {}
        if tok_emb_path and os.path.exists(tok_emb_path):
            try:
                self._load_tok_emb(tok_emb_path)
            except Exception as e:
                self.load_error = f"{type(e).__name__}: {e}"

    @staticmethod
    def _find_emb_and_vocab(d: dict):
        """dict(중첩 포함)에서 임베딩 행렬과 word2idx를 찾아낸다."""
        # 1) 직접 키
        emb = None
        for k in EmbeddingProvider.EMB_KEYS:
            if k in d and d[k] is not None:
                emb = np.asarray(d[k])
                if emb.ndim == 2:
                    break
                emb = None
        w2i = d.get("word2idx") or d.get("vocab") or d.get("stoi")
        if emb is not None and w2i is not None:
            return emb, w2i
        # 2) idx2word만 있으면 뒤집어서 word2idx 생성
        if emb is not None and "idx2word" in d:
            i2w = d["idx2word"]
            if isinstance(i2w, dict):
                w2i = {v: int(k) for k, v in i2w.items()}
            else:  # list
                w2i = {w: i for i, w in enumerate(i2w)}
            return emb, w2i
        # 3) 한 단계 중첩 탐색 (예: {"engine": {...}} / {"som": ..., "tok": {...}})
        for v in d.values():
            if isinstance(v, dict):
                e, w = EmbeddingProvider._find_emb_and_vocab(v)
                if e is not None and w is not None:
                    return e, w
        return None, None

    def _load_tok_emb(self, path: str):
        with open(path, "rb") as f:
            d = pickle.load(f)

        # 형식 4: 단어->벡터 딕셔너리
        if isinstance(d, dict) and d and all(
            isinstance(k, str) for k in list(d.keys())[:20]
        ) and all(
            isinstance(v, (list, np.ndarray)) for v in list(d.values())[:5]
        ) and not any(k in d for k in ("word2idx", "tok_emb", "idx2word", "vocab")):
            words = list(d.keys())
            mat = np.asarray([np.asarray(d[w], dtype=np.float64) for w in words])
            self.word2idx = {w: i for i, w in enumerate(words)}
            self.tok_emb = mat
            self.dim = mat.shape[1]
            self.vocab_size = len(words)
            self.mode = "tok_emb"
            return

        if not isinstance(d, dict):
            self.load_error = "pkl 최상위가 dict가 아님 — 인식 불가"
            return

        emb, w2i = self._find_emb_and_vocab(d)
        if emb is None or w2i is None:
            self.load_error = ("tok_emb/word2idx를 찾지 못함. "
                               f"최상위 키: {list(d.keys())[:8]}")
            return

        self.tok_emb = np.asarray(emb, dtype=np.float64)
        # 정수 인덱스 보장
        self.word2idx = {str(k): int(v) for k, v in w2i.items()}
        self.dim = self.tok_emb.shape[1]
        self.vocab_size = self.tok_emb.shape[0]
        self.mode = "tok_emb"

    def _word_vec(self, w: str) -> np.ndarray:
        if self.mode == "tok_emb":
            idx = self.word2idx.get(w)
            if idx is not None and 0 <= idx < self.tok_emb.shape[0]:
                return self.tok_emb[idx]
            return np.zeros(self.dim)
        # 해시 폴백
        if w not in self._cache:
            h = int(hashlib.md5(w.encode("utf-8")).hexdigest(), 16)
            r = np.random.default_rng(h % (2**32))
            self._cache[w] = r.normal(size=self.dim)
        return self._cache[w]

    def encode(self, sentence: str) -> np.ndarray:
        words = sentence.split()
        if not words:
            return np.zeros(self.dim)
        vs = [self._word_vec(w) for w in words]
        return np.mean(vs, axis=0)

    def encode_many(self, sentences) -> np.ndarray:
        X = np.stack([self.encode(s) for s in sentences])
        X = (X - X.mean(0)) / (X.std(0) + 1e-8)
        return X


# ======================================================================
# INT8 양자화
# ======================================================================
class Quantizer:
    def __init__(self, W):
        self.absmax = np.abs(W).max(axis=0) + 1e-8
        self.scale = self.absmax / 127.0

    def q(self, W):
        return np.clip(np.round(W / self.scale), -127, 127).astype(np.int8)

    def dq(self, Wq):
        return Wq.astype(np.float64) * self.scale


# ======================================================================
# Growing SOM
# ======================================================================
class GrowingSOM:
    def __init__(self, dim, init_nodes=16, grow_threshold=2.0, seed=0):
        rng = np.random.default_rng(seed)
        self.dim = dim
        self.W = rng.normal(scale=0.3, size=(init_nodes, dim))
        self.coords = rng.normal(scale=1.0, size=(init_nodes, 2))
        self.err = np.zeros(init_nodes)
        self.grow_threshold = grow_threshold
        self.round = 0

    @property
    def n(self):
        return self.W.shape[0]

    def bmu(self, x):
        d2 = np.einsum("nd,nd->n", self.W - x, self.W - x)
        i = int(np.argmin(d2))
        return i, float(np.sqrt(d2[i]))

    def train_step(self, X, lr, radius):
        for x in X:
            i, dist = self.bmu(x)
            self.err[i] += dist
            cd2 = np.einsum("nd,nd->n", self.coords - self.coords[i],
                            self.coords - self.coords[i])
            h = np.exp(-cd2 / (2 * radius ** 2))
            self.W += (lr * h)[:, None] * (x - self.W)

    def grow(self, max_add=4):
        added = 0
        order = np.argsort(self.err)[::-1]
        for i in order[:max_add]:
            if self.err[i] < self.grow_threshold:
                break
            new_w = self.W[i] + np.random.default_rng(self.n).normal(scale=0.1, size=self.dim)
            new_c = self.coords[i] + np.random.default_rng(self.n + 1).normal(scale=0.3, size=2)
            self.W = np.vstack([self.W, new_w])
            self.coords = np.vstack([self.coords, new_c])
            self.err = np.append(self.err, 0.0)
            self.err[i] *= 0.5
            added += 1
        return added

    # ---- 추론경로 (의미 가중 최단경로) ----
    def _knn_graph(self, k=6):
        D = np.linalg.norm(self.coords[:, None] - self.coords[None, :], axis=2)
        nbr = np.argsort(D, axis=1)[:, 1:k+1]
        return nbr

    def semantic_path(self, src, dst, k=6):
        import heapq
        if src == dst:
            return [src], 0.0, 0
        nbr = self._knn_graph(k)
        N = self.n
        dist = [float("inf")] * N
        prev = [-1] * N
        dist[src] = 0.0
        pq = [(0.0, src)]
        while pq:
            d, u = heapq.heappop(pq)
            if d > dist[u]:
                continue
            if u == dst:
                break
            for v in nbr[u]:
                w = np.linalg.norm(self.W[u] - self.W[v])
                nd = d + w
                if nd < dist[int(v)]:
                    dist[int(v)] = nd
                    prev[int(v)] = u
                    heapq.heappush(pq, (nd, int(v)))
        if not np.isfinite(dist[dst]):
            return [src, dst], float("inf"), -1
        path, cur = [], dst
        while cur != -1:
            path.append(cur); cur = prev[cur]
        path.reverse()
        return path, dist[dst], len(path) - 1


# ======================================================================
# 계측
# ======================================================================
def measure(gsom: GrowingSOM, X, tokens):
    bmus, qes, topo_fail = [], [], 0
    for x in X:
        d2 = np.einsum("nd,nd->n", gsom.W - x, gsom.W - x)
        order = np.argsort(d2)[:2]
        bmus.append(int(order[0]))
        qes.append(float(np.sqrt(d2[order[0]])))
        # 좌표상 1·2등 인접 여부
        c1, c2 = gsom.coords[order[0]], gsom.coords[order[1]]
        if np.linalg.norm(c1 - c2) > 1.5:
            topo_fail += 1
    occ = len(set(bmus))
    all_tok = [t for seq in tokens for t in seq]
    vocab = len(set(all_tok)) / len(all_tok) if all_tok else 0.0
    return {
        "nodes": gsom.n,
        "occupancy": occ,
        "occ_ratio": occ / gsom.n,
        "mean_qe": float(np.mean(qes)) if qes else 0.0,
        "topo_error": topo_fail / len(X) if len(X) else 0.0,
        "vocab_div": vocab,
    }


def collapse_warning(history, window=3):
    if len(history) < window + 1:
        return None
    rec = history[-(window+1):]
    def down(key):
        s = [h[key] for h in rec]
        return all(b <= a for a, b in zip(s[:-1], s[1:])) and s[0] > s[-1]
    if down("occupancy") and down("vocab_div") and down("mean_qe"):
        return "붕괴 경보: 점유율·어휘·QE 동반 감소 (자기출력 메아리방)"
    if down("occupancy") and rec[-1]["occ_ratio"] < 0.1:
        return "붕괴 경보: 점유율 10%↓ 한 점 수렴"
    return None


# ======================================================================
# 검색·크롤러 (네트워크 필요 — 로컬 실행 시 작동)
# ======================================================================
def _clean_text(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _split_sentences(text: str, min_len=12, max_len=220):
    """한국어/영어가 섞인 웹문서를 학습용 문장으로 쪼갠다."""
    text = _clean_text(text)
    # 문장부호 또는 줄바꿈성 구분자를 넓게 사용
    raw = re.split(r"(?<=[.!?。！？다요죠함음임됨됨니다])\s+|[\n\r]+|[•·]", text)
    out = []
    junk = ("cookie", "javascript", "copyright", "로그인", "회원가입", "개인정보", "구독", "광고")
    for s in raw:
        s = _clean_text(s)
        if not (min_len <= len(s) <= max_len):
            continue
        if any(j.lower() in s.lower() for j in junk):
            continue
        # 너무 URL/기호 위주인 문장 제거
        alpha_ko = len(re.findall(r"[A-Za-z가-힣0-9]", s))
        if alpha_ko / max(1, len(s)) < 0.45:
            continue
        out.append(s)
    return out


# ======================================================================
# 로컬 파일(docx / pdf / txt)에서 코퍼스 추출
# ======================================================================
def extract_text_from_docx(path: str) -> str:
    """Word(.docx) 본문 + 표 셀 텍스트를 추출."""
    try:
        import docx
    except Exception:
        raise RuntimeError("python-docx 미설치: pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # 표 안의 텍스트도 수집
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def extract_text_from_pdf(path: str) -> str:
    """PDF 본문 추출. pdfplumber 우선(레이아웃 양호), 실패 시 pypdf."""
    text = ""
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text() or ""
                if t.strip():
                    pages.append(t)
        text = "\n".join(pages)
    except Exception:
        text = ""
    if not text.strip():
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
        except Exception as e:
            raise RuntimeError(f"PDF 추출 실패: {e}")
    return text


def extract_corpus_from_file(path: str, filename: str | None = None,
                             min_len=12, max_len=220):
    """
    업로드 파일 경로에서 학습용 문장 리스트를 추출한다.
    지원: .docx .pdf .txt .md
    반환: (sentences, info) — info는 진단 문자열
    """
    name = (filename or path).lower()
    try:
        if name.endswith(".docx"):
            raw = extract_text_from_docx(path)
            kind = "docx"
        elif name.endswith(".pdf"):
            raw = extract_text_from_pdf(path)
            kind = "pdf"
        elif name.endswith((".txt", ".md")):
            with open(path, "rb") as f:
                data = f.read()
            raw = None
            for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
                try:
                    raw = data.decode(enc); break
                except Exception:
                    continue
            raw = raw or data.decode("utf-8", "ignore")
            kind = "txt"
        elif name.endswith(".doc"):
            return [], "구버전 .doc는 미지원입니다. .docx로 저장 후 올려주세요."
        else:
            return [], f"지원하지 않는 형식: {name.split('.')[-1]}"
    except Exception as e:
        return [], f"{kind if 'kind' in dir() else '파일'} 추출 오류: {e}"

    sents = _split_sentences(raw, min_len=min_len, max_len=max_len)
    # 중복 제거(순서 보존)
    seen, uniq = set(), []
    for s in sents:
        if s not in seen:
            seen.add(s); uniq.append(s)
    info = f"{kind} · 원문 {len(raw)}자 → 문장 {len(uniq)}개"
    return uniq, info


_BROWSER_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/124.0.0.0 Safari/537.36"),
    "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
               "image/avif,image/webp,*/*;q=0.8"),
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
    "Accept-Encoding": "identity",
    "Referer": "https://www.google.com/",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
}


def _fetch_url_text(url: str, timeout=12) -> str:
    import urllib.request
    req = urllib.request.Request(url, headers=_BROWSER_HEADERS)
    with urllib.request.urlopen(req, timeout=timeout) as r:
        raw = r.read()
        ctype = r.headers.get("Content-Type", "")
    # 간단 디코딩. 한국어 페이지는 utf-8/euc-kr 섞일 수 있음.
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return raw.decode(enc, "ignore")
        except Exception:
            pass
    return raw.decode("utf-8", "ignore")


class _TextExtractorHTML:
    """html.parser를 함수 내부 import 없이 쓰기 위한 가벼운 본문 추출기."""
    pass


def _html_to_text(html: str) -> str:
    from html.parser import HTMLParser

    class TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.skip_depth = 0
            self.parts = []
        def handle_starttag(self, tag, attrs):
            if tag.lower() in ("script", "style", "noscript", "svg", "canvas", "header", "footer", "nav"):
                self.skip_depth += 1
        def handle_endtag(self, tag):
            if tag.lower() in ("script", "style", "noscript", "svg", "canvas", "header", "footer", "nav") and self.skip_depth:
                self.skip_depth -= 1
        def handle_data(self, data):
            if self.skip_depth == 0:
                t = data.strip()
                if t:
                    self.parts.append(t)

    ex = TextExtractor()
    ex.feed(html)
    return " ".join(ex.parts)


def _extract_links_generic(html: str, max_results: int):
    """검색결과 HTML에서 외부 URL 추출 (DDG 리다이렉트 디코딩 포함)."""
    import urllib.parse
    links = []
    for m in re.findall(r'href=["\']([^"\']+)["\']', html):
        href = m.replace("&amp;", "&")
        real = None
        if "uddg=" in href:  # DuckDuckGo 리다이렉트
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(href).query)
            if "uddg" in qs:
                real = urllib.parse.unquote(qs["uddg"][0])
        elif "/url?q=" in href:  # Google 리다이렉트
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(href).query)
            if "q" in qs:
                real = qs["q"][0]
        elif href.startswith("http://") or href.startswith("https://"):
            real = href
        if not real:
            continue
        if any(b in real for b in ["duckduckgo.com", "bing.com", "google.com",
                                   "javascript:", "mailto:", "microsoft.com/",
                                   "go.microsoft.com", "w3.org"]):
            continue
        if real not in links:
            links.append(real)
        if len(links) >= max_results:
            break
    return links


def _search_links(query: str, max_results=8):
    """
    여러 검색 엔진을 순서대로 시도. 하나가 막히면 다음으로 폴백.
    반환: (links, log) — log는 각 엔진 시도 결과(디버그/화면표시용).
    """
    import urllib.parse
    q = urllib.parse.quote(query)
    engines = [
        ("DuckDuckGo Lite", f"https://lite.duckduckgo.com/lite/?q={q}"),
        ("DuckDuckGo HTML", f"https://html.duckduckgo.com/html/?q={q}"),
        ("Bing", f"https://www.bing.com/search?q={q}&setlang=ko"),
        ("Mojeek", f"https://www.mojeek.com/search?q={q}"),
    ]
    log = []
    for name, url in engines:
        try:
            html = _fetch_url_text(url)
            links = _extract_links_generic(html, max_results)
            log.append(f"{name}: {len(links)}개 링크")
            if links:
                return links, log
        except Exception as e:
            log.append(f"{name}: 실패({type(e).__name__} {e})")
            continue
    return [], log


def _duckduckgo_links(query: str, max_results=8):
    """하위호환 래퍼."""
    links, _ = _search_links(query, max_results)
    return links


def crawl_urls(urls, max_sentences=300, delay=0.3):
    """URL 목록을 직접 크롤링해서 문장 리스트와 소스 로그를 반환한다."""
    sentences, sources = [], []
    for url in urls:
        try:
            html = _fetch_url_text(url)
            text = _html_to_text(html)
            ss = _split_sentences(text)
            if ss:
                sources.append({"url": url, "sentences": len(ss)})
                sentences.extend(ss)
            if delay:
                time.sleep(delay)
        except Exception as e:
            sources.append({"url": url, "error": f"{type(e).__name__}: {e}"})
            continue
    # 중복 제거
    seen, out = set(), []
    for s in sentences:
        key = s.lower().strip()
        if key not in seen:
            seen.add(key); out.append(s)
        if len(out) >= max_sentences:
            break
    return out, sources


def crawl_topic(topic: str, max_pages=5, max_sentences=300, extra_urls=None, delay=0.3, return_sources=False):
    """
    '이 분야 검색해서 학습' 명령 처리.

    - 여러 검색 엔진(DDG Lite/HTML, Bing, Mojeek)을 폴백하며 URL 수집
    - 사용자가 직접 넣은 URL도 함께 크롤링
    - 본문 문장 추출 후 중복 제거
    - 검색/크롤링 실패 시 sources에 진단 로그를 담아 반환

    return_sources=True면 (sentences, sources, links)를 반환한다.
    """
    if not topic and not extra_urls:
        return ([], ["입력된 검색어/URL 없음"], []) if return_sources else []

    urls = []
    search_log = []
    if topic and topic.strip():
        found, search_log = _search_links(topic.strip(), max_results=max_pages)
        urls.extend(found)
    if extra_urls:
        for u in extra_urls:
            u = u.strip()
            if u and (u.startswith("http://") or u.startswith("https://")) and u not in urls:
                urls.append(u)

    if not urls:
        diag = ["검색 결과 0건 — 엔진별 시도:"] + search_log
        diag.append("→ 모든 검색 엔진이 차단되었거나 네트워크가 막혀 있습니다. "
                    "URL 직접 입력을 사용해 보세요.")
        return ([], diag, []) if return_sources else []

    sentences, sources = crawl_urls(urls[:max_pages], max_sentences=max_sentences, delay=delay)
    if return_sources:
        return sentences, sources + search_log, urls[:max_pages]
    return sentences


# ======================================================================
# LLM 브리지 (답변 루프 — API 키 필요)
# ======================================================================
def llm_answer(
    question: str,
    context_sentences,
    model="gpt-4o-mini",
    temperature=0.3,
    max_tokens=400,
    api_key: str | None = None,
    base_url: str | None = None,
    system_prompt: str | None = None,
):
    """
    학습된 SOM 코퍼스에서 뽑은 맥락으로 OpenAI-compatible LLM 답변 생성.

    - api_key가 전달되면 우선 사용하고, 없으면 OPENAI_API_KEY 환경변수를 사용합니다.
    - base_url을 전달하면 OpenAI 호환 서버(Ollama proxy, vLLM, LM Studio 등)에도 연결할 수 있습니다.
      예: https://api.openai.com/v1 또는 http://localhost:1234/v1
    """
    try:
        from openai import OpenAI
    except Exception:
        return "[openai 미설치] pip install openai 후 사용"

    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        return "[API KEY 미설정] 사이드바에 API Key를 입력하거나 OPENAI_API_KEY 환경변수를 설정하세요."

    kwargs = {"api_key": key}
    if base_url:
        kwargs["base_url"] = base_url.rstrip("/")
    client = OpenAI(**kwargs)

    ctx = "\n".join(f"- {s}" for s in context_sentences[:15])
    user = f"[학습된 맥락]\n{ctx}\n\n[질문]\n{question}"

    # 시스템 프롬프트 미사용: 맥락만 user 메시지로 전달.
    # system_prompt를 명시적으로 넘긴 경우에만 system 메시지를 추가한다.
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": user})

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"[LLM 호출 실패] {type(e).__name__}: {e}"


# ======================================================================
# 전체 상태 + pkl 누적 저장/불러오기
# ======================================================================
class SelfLoopState:
    def __init__(self, dim=64):
        self.dim = dim
        self.gsom = GrowingSOM(dim=dim)
        self.corpus: list[str] = []       # 학습된 모든 문장(코퍼스)
        self.history: list[dict] = []     # 라운드별 계측
        self.created = time.time()

    def save(self, path):
        # SOM 가중치는 INT8 양자화해서 저장 (용량 보완)
        quant = Quantizer(self.gsom.W)
        blob = {
            "dim": self.dim,
            "W_int8": quant.q(self.gsom.W),
            "W_scale": quant.scale,
            "coords": self.gsom.coords,
            "err": self.gsom.err,
            "round": self.gsom.round,
            "corpus": self.corpus,
            "history": self.history,
            "created": self.created,
        }
        with open(path, "wb") as f:
            pickle.dump(blob, f)

    @classmethod
    def load(cls, path):
        with open(path, "rb") as f:
            b = pickle.load(f)
        st = cls(dim=b["dim"])
        # 역양자화로 SOM 복원
        st.gsom.W = b["W_int8"].astype(np.float64) * b["W_scale"]
        st.gsom.coords = b["coords"]
        st.gsom.err = b["err"]
        st.gsom.round = b.get("round", 0)
        st.corpus = b["corpus"]
        st.history = b["history"]
        st.created = b.get("created", time.time())
        return st
